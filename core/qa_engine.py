# core/qa_engine.py
"""
QA Engine для работы с документацией.
Поддерживает:
- Семантический поиск по документам
- Гибридный поиск (FAISS + BM25)
- Извлечение таблиц и формул
- Индексацию документов
- Кэширование индекса
- Поиск определений
- Опциональный LLM (OpenAI) для генерации ответов
"""

import hashlib
import os
import pickle
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, cast

# noinspection PyUnresolvedReferences
import faiss

# noinspection PyUnresolvedReferences
import numpy as np

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None  # type: ignore

try:
    from rank_bm25 import BM25Okapi
except ImportError:
    BM25Okapi = None  # type: ignore

try:
    from huggingface_hub import snapshot_download
except ImportError:
    snapshot_download = None  # type: ignore

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None  # type: ignore


class QASystem:
    """Система вопрос-ответ на основе документов с поддержкой LLM."""

    def __init__(self, use_llm: bool = False) -> None:
        if SentenceTransformer is None:
            raise ImportError("Не установлен sentence-transformers")

        print("🔄 Загрузка модели эмбеддингов...")
        model_cls = cast(Any, SentenceTransformer)
        # noinspection SpellCheckingInspection
        self.model = model_cls("intfloat/multilingual-e5-small")
        self.dimension = 384

        # noinspection PyTypeChecker
        self.index = None
        self.chunks: List[Dict[str, Any]] = []
        self.bm25_index = None
        self.is_ready = False

        self.use_llm = use_llm
        self.llm_engine = None

        if self.use_llm:
            if OpenAI is None:
                raise ImportError("openai не установлен. Установи: pip install openai")

            api_key = os.getenv("OPENAI_API_KEY", "").strip()
            if not api_key:
                raise ValueError("OPENAI_API_KEY не найден в переменных окружения")

            openai_cls = cast(Any, OpenAI)
            self.llm_engine = openai_cls(api_key=api_key)
            print("✅ OpenAI клиент инициализирован")

        self.stop_phrases = [
            ".", ",", "q", "t", "r", "tv", "tn", "tot", "zot",
            "δ", "λ", "Q", "R", "A", "L", "("
        ]
        self.definitions_cache: Dict[str, Dict[str, Any]] = {}
        self._chunk_size = 1000
        self._chunk_overlap = 200

    @staticmethod
    def normalize_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def is_bad_chunk(self, text: str) -> bool:
        text_stripped = self.normalize_text(text).lower()
        if len(text_stripped) < 50:
            return True
        if text_stripped in self.stop_phrases:
            return True
        if sum(ch.isalpha() for ch in text_stripped) < 20:
            return True
        return False

    @staticmethod
    def chunk_text(
        text: str,
        doc_name: str,
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> List[Dict[str, Any]]:
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            return []

        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks: List[Dict[str, Any]] = []
        current = ""
        chunk_id = 0

        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue

            if len(current) + len(sent) + 1 <= chunk_size:
                current = f"{current} {sent}".strip()
            else:
                if current:
                    chunks.append({
                        "id": f"{doc_name}_{chunk_id}",
                        "text": current,
                        "doc_name": doc_name,
                        "chunk_id": chunk_id,
                    })
                    chunk_id += 1

                    if overlap > 0:
                        tail = current[-overlap:]
                        current = f"{tail} {sent}".strip()
                    else:
                        current = sent

        if current:
            chunks.append({
                "id": f"{doc_name}_{chunk_id}",
                "text": current,
                "doc_name": doc_name,
                "chunk_id": chunk_id,
            })

        return chunks

    @staticmethod
    def read_docx(file_path: str) -> str:
        try:
            # noinspection PyUnresolvedReferences
            from docx import Document
        except ImportError:
            return ""

        try:
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]

            for table_idx, table in enumerate(doc.tables):
                table_lines = [f"Таблица {table_idx + 1}"]
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        table_lines.append(" | ".join(cells))
                if len(table_lines) > 1:
                    parts.append("\n".join(table_lines))

            return "\n".join(parts)
        except (OSError, ValueError):
            return ""

    @staticmethod
    def read_pdf(file_path: str) -> str:
        # noinspection PyUnresolvedReferences
        import fitz  # type: ignore

        try:
            parts: List[str] = []
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                parts.append(f"--- page {page_num + 1} ---")
                parts.append(page.get_text())

            try:
                # noinspection PyUnresolvedReferences
                import pdfplumber
            except ImportError:
                pdfplumber = None  # type: ignore

            if pdfplumber is not None:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            tables = page.extract_tables() or []
                            for table_idx, table in enumerate(tables):
                                lines = [f"Таблица {table_idx + 1}, страница {page_num + 1}"]
                                for row in table:
                                    if row:
                                        row_text = [str(cell or "").strip() for cell in row]
                                        if any(row_text):
                                            lines.append(" | ".join(row_text))
                                if len(lines) > 1:
                                    parts.append("\n".join(lines))
                except (OSError, ValueError):
                    pass

            return "\n".join(parts)
        except (OSError, ValueError):
            return ""

    @staticmethod
    def read_rtf(file_path: str) -> str:
        try:
            # noinspection PyUnresolvedReferences
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            print("⚠️ striprtf не установлен. pip install striprtf")
            return ""

        try:
            raw_text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
            return rtf_to_text(raw_text).strip()
        except (OSError, ValueError):
            return ""

    def read_file(self, file_path: Path) -> str:
        ext = file_path.suffix.lower()

        try:
            if ext in {".docx", ".doc"}:
                return self.read_docx(str(file_path))
            if ext == ".pdf":
                return self.read_pdf(str(file_path))
            if ext == ".rtf":
                return self.read_rtf(str(file_path))
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except (OSError, ValueError):
            return ""

    @staticmethod
    def get_documents_hash(documents_dir: Path) -> str:
        md5 = hashlib.md5()
        for file_path in sorted(documents_dir.glob("*")):
            if file_path.is_file() and file_path.suffix.lower() in {
                ".docx", ".doc", ".pdf", ".rtf", ".txt", ".md", ".csv", ".json", ".xml"
            }:
                try:
                    with open(file_path, "rb") as f:
                        for chunk in iter(lambda: f.read(4096), b""):
                            md5.update(chunk)
                except OSError:
                    continue
        return md5.hexdigest()

    @staticmethod
    def get_hash_file_path(documents_dir: Path) -> Path:
        return documents_dir.parent / "processed" / "docs_hash.txt"

    def load_saved_hash(self, documents_dir: Path) -> Optional[str]:
        hash_file = self.get_hash_file_path(documents_dir)
        if not hash_file.exists():
            return None
        try:
            return hash_file.read_text(encoding="utf-8").strip()
        except (OSError, ValueError):
            return None

    def save_hash(self, documents_dir: Path, hash_value: str) -> None:
        hash_file = self.get_hash_file_path(documents_dir)
        hash_file.parent.mkdir(parents=True, exist_ok=True)
        hash_file.write_text(hash_value, encoding="utf-8")

    @staticmethod
    def tokenize(text: str) -> List[str]:
        return re.findall(r"[A-Za-zА-Яа-я0-9\-]{2,}", text.lower())

    def build_bm25_index(self) -> None:
        if not self.chunks or BM25Okapi is None:
            self.bm25_index = None
            return

        bm25_cls = cast(Any, BM25Okapi)
        tokenized = [self.tokenize(chunk["text"]) for chunk in self.chunks]
        self.bm25_index = bm25_cls(tokenized)

    def save_index(self, index_dir: Path) -> None:
        if self.index is None:
            return

        try:
            index_dir.mkdir(parents=True, exist_ok=True)
            # noinspection PyUnresolvedReferences
            faiss.write_index(self.index, str(index_dir / "qa_index.faiss"))
            with open(index_dir / "chunks.pkl", "wb") as f:
                pickle.dump(self.chunks, f)
        except (OSError, ValueError):
            return

    def load_index(self, index_dir: Path) -> bool:
        index_path = index_dir / "qa_index.faiss"
        chunks_path = index_dir / "chunks.pkl"

        if not index_path.exists() or not chunks_path.exists():
            return False

        try:
            # noinspection PyUnresolvedReferences
            self.index = faiss.read_index(str(index_path))
            with open(chunks_path, "rb") as f:
                self.chunks = pickle.load(f)
            self.build_bm25_index()
            self.is_ready = True
            return True
        except (OSError, ValueError, pickle.UnpicklingError):
            return False

    def index_documents(self, documents_dir: Path) -> bool:
        documents_dir.mkdir(parents=True, exist_ok=True)
        index_dir = documents_dir.parent / "processed"

        current_hash = self.get_documents_hash(documents_dir)
        saved_hash = self.load_saved_hash(documents_dir)

        if saved_hash == current_hash and self.load_index(index_dir):
            return True

        if snapshot_download is not None:
            try:
                snapshot_fn = cast(Any, snapshot_download)
                snapshot_fn(
                    repo_id="Lana49/engineering-docs",
                    repo_type="dataset",
                    local_dir=str(documents_dir),
                    allow_patterns=["*.docx", "*.doc", "*.pdf", "*.rtf", "*.txt", "*.md", "*.csv", "*.json", "*.xml"],
                    local_dir_use_symlinks=False,
                    force_download=False,
                )
            except (OSError, ValueError):
                pass

        all_chunks: List[Dict[str, Any]] = []

        for file_path in documents_dir.glob("*"):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in {
                ".docx", ".doc", ".pdf", ".rtf", ".txt", ".md", ".csv", ".json", ".xml"
            }:
                continue

            text = self.read_file(file_path)
            if not text.strip():
                continue

            chunks = self.chunk_text(text, file_path.stem, self._chunk_size, self._chunk_overlap)
            all_chunks.extend(chunks)

        all_chunks = [c for c in all_chunks if not self.is_bad_chunk(c["text"])]
        if not all_chunks:
            self.is_ready = False
            return False

        texts = [f"passage: {c['text']}" for c in all_chunks]
        embeddings = self.model.encode(texts, show_progress_bar=True)
        # noinspection PyUnresolvedReferences,PyTypeChecker
        embeddings = np.asarray(embeddings, dtype=np.float32)

        # noinspection PyUnresolvedReferences
        norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        embeddings = embeddings / norms

        # noinspection PyUnresolvedReferences
        self.index = faiss.IndexFlatIP(self.dimension)
        self.index.add(embeddings)

        self.chunks = all_chunks
        self.build_bm25_index()
        self.save_index(index_dir)
        self.save_hash(documents_dir, current_hash)

        self.is_ready = True
        return True

    def search_bm25(self, query: str, top_k: int = 10) -> List[Dict[str, Any]]:
        if self.bm25_index is None:
            return []

        try:
            tokenized_query = self.tokenize(query)
            scores = self.bm25_index.get_scores(tokenized_query)
            # noinspection PyUnresolvedReferences,PyTypeChecker
            top_indices = np.argsort(scores)[::-1][:top_k]
            results: List[Dict[str, Any]] = []

            for idx in top_indices:
                idx_int = int(idx)
                if idx_int < len(self.chunks) and float(scores[idx_int]) > 0:
                    results.append({
                        "text": self.chunks[idx_int]["text"],
                        "doc_name": self.chunks[idx_int]["doc_name"],
                        "score": float(scores[idx_int]),
                        "idx": idx_int,
                        "search_type": "bm25",
                    })
            return results
        except ValueError:
            return []

    def search(self, query: str, top_k: int = 8) -> List[Dict[str, Any]]:
        if not self.is_ready or self.index is None:
            return []

        # noinspection PyTypeChecker
        query_emb = self.model.encode([f"query: {query}"])
        # noinspection PyUnresolvedReferences,PyTypeChecker
        query_emb = np.asarray(query_emb, dtype=np.float32)

        # noinspection PyUnresolvedReferences
        norms = np.linalg.norm(query_emb, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        query_emb = query_emb / norms

        # noinspection PyUnresolvedReferences
        scores, indices = self.index.search(query_emb, max(top_k * 2, 10))
        faiss_results: List[Dict[str, Any]] = []

        for score, idx in zip(scores[0], indices[0]):
            idx_int = int(idx)
            if 0 <= idx_int < len(self.chunks) and float(score) > 0.25:
                faiss_results.append({
                    "text": self.chunks[idx_int]["text"],
                    "doc_name": self.chunks[idx_int]["doc_name"],
                    "score": float(score),
                    "idx": idx_int,
                    "search_type": "faiss",
                })

        bm25_results = self.search_bm25(query, top_k=top_k)

        combined: Dict[str, Dict[str, Any]] = {}
        for item in faiss_results + bm25_results:
            key = f"{item['doc_name']}::{item['idx']}"
            if key not in combined:
                combined[key] = item
            else:
                combined[key]["score"] = max(combined[key]["score"], item["score"])

        results = sorted(combined.values(), key=lambda x: x["score"], reverse=True)
        return results[:top_k]

    @staticmethod
    def extract_tables(text: str) -> List[str]:
        tables: List[str] = []
        pattern = r"(Таблица.*?(?:\n.+)+)"
        matches = re.findall(pattern, text, flags=re.MULTILINE)
        tables.extend(matches)
        return tables

    @staticmethod
    def extract_formulas(text: str) -> List[Dict[str, Any]]:
        formulas: List[Dict[str, Any]] = []
        patterns = [
            r"[A-Za-zА-Яа-я][\w]*\s*=\s*[^.\n]+",
            r"[^.\n]*=\s*[^.\n]*",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                match = match.strip()
                if len(match) < 3 or "=" not in match:
                    continue

                formulas.append({
                    "raw": match,
                    "variables": list(set(re.findall(r"[A-Za-zА-Яа-я][_\w]*", match))),
                    "has_operator": True,
                })

        unique_formulas: List[Dict[str, Any]] = []
        seen = set()
        for formula in formulas:
            raw = formula["raw"]
            if raw not in seen:
                seen.add(raw)
                unique_formulas.append(formula)

        return unique_formulas

    def search_with_formulas(self, query: str, top_k: int = 5) -> Dict[str, Any]:
        results = self.search(query, top_k=top_k)

        all_tables: List[str] = []
        all_formulas: List[Dict[str, Any]] = []

        for item in results:
            text = item.get("text", "")
            all_tables.extend(self.extract_tables(text))
            all_formulas.extend(self.extract_formulas(text))

        unique_tables: List[str] = []
        seen_tables = set()
        for table in all_tables:
            if table not in seen_tables:
                seen_tables.add(table)
                unique_tables.append(table)

        unique_formulas: List[Dict[str, Any]] = []
        seen_formulas = set()
        for formula in all_formulas:
            raw = formula["raw"]
            if raw not in seen_formulas:
                seen_formulas.add(raw)
                unique_formulas.append(formula)

        return {
            "results": results,
            "tables": unique_tables,
            "formulas": unique_formulas,
        }

    @staticmethod
    def _build_llm_prompt(
        question: str,
        chunks: List[Dict[str, Any]],
        tables: List[str],
        formulas: List[Dict[str, Any]],
    ) -> str:
        context_parts: List[str] = []

        for i, chunk in enumerate(chunks[:5], 1):
            context_parts.append(
                f"[Фрагмент {i} | Источник: {chunk['doc_name']}]\n{chunk['text']}"
            )

        if tables:
            context_parts.append("\n[Таблицы]")
            for table in tables[:3]:
                context_parts.append(table)

        if formulas:
            context_parts.append("\n[Формулы]")
            for formula in formulas[:5]:
                context_parts.append(formula["raw"])

        context = "\n\n".join(context_parts)

        return f"""
Ты отвечаешь только на основе контекста из технической документации ниже.
Если точного ответа в контексте нет, честно скажи об этом.
Не выдумывай факты.
Отвечай кратко, по делу и на русском языке.

Вопрос:
{question}

Контекст:
{context}
""".strip()

    def answer(self, question: str, top_k: int = 5) -> Dict[str, Any]:
        search_results = self.search_with_formulas(question, top_k)
        relevant = search_results["results"]

        if not relevant:
            return {
                "question": question,
                "answer": "❌ Информация по вашему вопросу не найдена в документации.",
                "sources": [],
                "tables": [],
                "formulas": [],
                "llm_used": False,
            }

        cleaned_chunks = [c for c in relevant if not self.is_bad_chunk(c["text"])]
        if not cleaned_chunks:
            cleaned_chunks = relevant[:2]

        all_tables = search_results.get("tables", [])
        all_formulas = search_results.get("formulas", [])
        llm_used = False

        if self.use_llm and self.llm_engine is not None:
            try:
                prompt = self._build_llm_prompt(
                    question=question,
                    chunks=cleaned_chunks,
                    tables=all_tables,
                    formulas=all_formulas,
                )

                response = self.llm_engine.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Ты — инженерный помощник по строительной документации."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.2,
                    max_tokens=1000,
                )
                answer_text = response.choices[0].message.content.strip()
                llm_used = True
            except (ValueError, RuntimeError):
                answer_text = cleaned_chunks[0]["text"].split(". ")[0].strip()
        else:
            answer_text = cleaned_chunks[0]["text"].split(". ")[0].strip()

        if not answer_text.endswith("."):
            answer_text += "."

        if all_tables:
            answer_text += "\n\n📊 Найденные таблицы:\n"
            for table in all_tables[:2]:
                answer_text += f"\n{table[:700]}\n"

        if all_formulas:
            answer_text += "\n\n📐 Найденные формулы:\n"
            for formula in all_formulas[:3]:
                raw = formula.get("raw", "")
                answer_text += f"\n{raw}\n"
                variables = formula.get("variables", [])
                if variables:
                    answer_text += f"Переменные: {', '.join(variables[:5])}\n"

        return {
            "question": question,
            "answer": answer_text,
            "sources": cleaned_chunks,
            "tables": all_tables[:3],
            "formulas": all_formulas[:5],
            "llm_used": llm_used,
        }

    @staticmethod
    def _init_definitions() -> Dict[str, str]:
        return {
            "теплопередача": "Процесс переноса тепла от более нагретого тела к менее нагретому.",
            "коэффициент теплопередачи": "Величина, характеризующая интенсивность передачи тепла через ограждающую конструкцию.",
            "сопротивление теплопередаче": "Способность конструкции препятствовать прохождению теплового потока.",
            "точка росы": "Температура, при которой водяной пар в воздухе начинает конденсироваться.",
            "влажность": "Содержание водяного пара в воздухе или материале.",
        }

    def find_definition(self, term: str) -> Dict[str, Any]:
        term_lower = term.lower().strip()

        if term_lower in self.definitions_cache:
            return self.definitions_cache[term_lower]

        all_definitions = self._init_definitions()

        if term_lower in all_definitions:
            result = {
                "term": term,
                "definition": all_definitions[term_lower],
                "source": "built-in",
                "found": True,
            }
            self.definitions_cache[term_lower] = result
            return result

        search_results = self.search(term, top_k=3)
        if search_results:
            result = {
                "term": term,
                "definition": search_results[0]["text"][:500],
                "source": search_results[0]["doc_name"],
                "found": True,
            }
            self.definitions_cache[term_lower] = result
            return result

        result = {
            "term": term,
            "definition": "Определение не найдено.",
            "source": None,
            "found": False,
        }
        self.definitions_cache[term_lower] = result
        return result