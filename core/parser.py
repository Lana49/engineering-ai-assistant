# -*- coding: utf-8 -*-
"""
Универсальный парсер документов для инженерной базы знаний.

Поддерживает:
- .pdf   -> PyMuPDF + pdfplumber + OCR (tesseract) как fallback
- .docx  -> python-docx
- .doc   -> textutil (macOS) / textract + antiword (Linux)
- .rtf   -> striprtf
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".rtf"}

IS_MACOS = sys.platform == "darwin"
IS_LINUX = sys.platform.startswith("linux")


def command_exists(name: str) -> bool:
    """Проверяет, доступна ли системная команда."""
    return shutil.which(name) is not None


def is_supported_file(path: Path) -> bool:
    """Проверяет, поддерживается ли файл по расширению."""
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


def safe_print(message: str) -> None:
    """Безопасный вывод в консоль."""
    try:
        print(message)
    except (OSError, UnicodeEncodeError):  # ← ИСПРАВЛЕНО: конкретные исключения
        pass


# =========================
# PDF
# =========================

def read_pdf_pymupdf(file_path: str | Path) -> tuple[str, str | None]:
    """Читает PDF через PyMuPDF."""
    try:
        import pymupdf
    except ImportError:
        try:
            import fitz as pymupdf  # type: ignore
        except ImportError:
            return "", "PyMuPDF не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        parts: list[str] = []
        with pymupdf.open(str(path)) as doc:  # type: ignore
            for page_num, page in enumerate(doc, start=1):
                try:
                    page_text = page.get_text("text") or ""
                    if page_text.strip():
                        parts.append(page_text.strip())
                except (AttributeError, ValueError, TypeError, RuntimeError) as exc:
                    safe_print(f"⚠️ PyMuPDF: ошибка страницы {page_num} в {path.name}: {exc}")

        text = "\n\n".join(parts).strip()
        if text:
            return text, "PyMuPDF"
        return "", "текст не извлечён (возможно, сканированный PDF)"
    except (OSError, ValueError, TypeError, RuntimeError) as exc:
        return "", f"PyMuPDF ошибка: {exc}"


def read_pdf_pdfplumber(file_path: str | Path) -> tuple[str, str | None]:
    """Читает PDF через pdfplumber как fallback."""
    try:
        import pdfplumber
    except ImportError:
        return "", "pdfplumber не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    text = page.extract_text() or ""
                    if text.strip():
                        parts.append(text.strip())
                except (AttributeError, ValueError, TypeError) as exc:
                    safe_print(f"⚠️ pdfplumber: ошибка страницы {page_num} в {path.name}: {exc}")

        text = "\n\n".join(parts).strip()
        if text:
            return text, "pdfplumber"
        return "", "текст не извлечён (возможно, сканированный PDF)"
    except (OSError, ValueError, TypeError, RuntimeError) as exc:
        return "", f"pdfplumber ошибка: {exc}"


def read_pdf_ocr(file_path: str | Path) -> tuple[str, str | None]:
    """OCR fallback для сканированных PDF через Tesseract."""
    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        import pymupdf
    except ImportError:
        try:
            import fitz as pymupdf  # type: ignore
        except ImportError:
            return "", "PyMuPDF не установлен (нужен для OCR)"

    try:
        import pytesseract  # noqa: F401
    except ImportError:
        return "", "pytesseract не установлен"

    try:
        from PIL import Image  # noqa: F401
    except ImportError:
        return "", "Pillow не установлен"

    if not command_exists("tesseract"):
        return "", "Tesseract не установлен в системе"

    try:
        parts: list[str] = []
        with pymupdf.open(str(path)) as doc:  # type: ignore
            total_pages = len(doc)
            safe_print(f"📄 OCR страниц {path.name}: {total_pages}")

            for page_num, page in enumerate(doc, start=1):
                try:
                    matrix = pymupdf.Matrix(2.5, 2.5)  # type: ignore
                    pix = page.get_pixmap(matrix=matrix)

                    if pix.alpha:
                        img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples).convert("RGB")
                    else:
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    text = pytesseract.image_to_string(
                        img,
                        lang="rus+eng",
                        config="--psm 6 --oem 3",
                    )
                    if text.strip():
                        parts.append(text.strip())
                        safe_print(f"   ✅ Страница {page_num}/{total_pages} распознана")
                    else:
                        safe_print(f"   ⚠️ Страница {page_num}/{total_pages} пустая")
                except (AttributeError, ValueError, TypeError, OSError, RuntimeError) as exc:
                    safe_print(f"   ⚠️ Страница {page_num}/{total_pages} ошибка: {exc}")

        text = "\n\n".join(parts).strip()
        if text:
            return text, "OCR (Tesseract)"
        return "", "OCR не распознал текст"
    except (OSError, ValueError, TypeError, RuntimeError) as exc:
        return "", f"OCR ошибка: {exc}"


def read_pdf(file_path: str | Path) -> tuple[str, str | None]:
    """Читает PDF с несколькими fallback-стратегиями."""
    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл не существует или пустой"

    safe_print(f"📄 Чтение PDF: {path.name}")

    for reader in (read_pdf_pymupdf, read_pdf_pdfplumber, read_pdf_ocr):
        text, source = reader(path)
        if text:
            return text, source

    return "", "все методы извлечения PDF не сработали"


# =========================
# DOCX
# =========================

def read_docx(file_path: str | Path) -> tuple[str, str | None]:
    """Читает DOCX через python-docx."""
    try:
        from docx import Document
    except ImportError:
        return "", "python-docx не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        doc = Document(str(path))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()
        if text:
            return text, "python-docx"
        return "", "текст не извлечён"
    except (OSError, ValueError, TypeError, ImportError) as exc:  # ← ИСПРАВЛЕНО
        return "", f"python-docx ошибка: {exc}"


# =========================
# DOC
# =========================

def read_doc_textutil(file_path: str | Path) -> tuple[str, str | None]:
    """Читает .doc через textutil на macOS."""
    if not IS_MACOS:
        return "", "textutil доступен только на macOS"

    if not command_exists("textutil"):
        return "", "textutil не найден в системе"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        result = subprocess.run(
            ["textutil", "-stdout", "-convert", "txt", str(path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip(), "textutil"
        stderr = (result.stderr or "").strip()
        return "", f"textutil не вернул текст{': ' + stderr if stderr else ''}"
    except (subprocess.TimeoutExpired, subprocess.SubprocessError, OSError) as exc:
        return "", f"textutil ошибка: {exc}"


def read_doc_textract(file_path: str | Path) -> tuple[str, str | None]:
    """Читает .doc через textract."""
    try:
        import textract  # noqa: F401
    except ImportError:
        return "", "textract не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        import textract
        data = textract.process(str(path))
        text = data.decode("utf-8", errors="ignore").strip()
        if text:
            return text, "textract"
        return "", "textract не вернул текст"
    except (OSError, ValueError, TypeError, RuntimeError) as exc:  # ← ИСПРАВЛЕНО
        return "", f"textract ошибка: {exc}"


def read_doc_antiword(file_path: str | Path) -> tuple[str, str | None]:
    """Читает .doc через antiword."""
    if not command_exists("antiword"):
        return "", "antiword не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip(), "antiword"
        stderr = (result.stderr or "").strip()
        return "", f"antiword не вернул текст{': ' + stderr if stderr else ''}"
    except (subprocess.TimeoutExpired, subprocess.SubprocessError, OSError) as exc:
        return "", f"antiword ошибка: {exc}"


def read_doc(file_path: str | Path) -> tuple[str, str | None]:
    """Читает .doc с платформенным fallback."""
    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл не существует или пустой"

    safe_print(f"📄 Чтение DOC: {path.name}")

    if IS_MACOS:
        text, source = read_doc_textutil(path)
        if text:
            return text, source

    if IS_LINUX:
        text, source = read_doc_textract(path)
        if text:
            return text, source

        text, source = read_doc_antiword(path)
        if text:
            return text, source

    # fallback для других платформ
    text, source = read_doc_textract(path)
    if text:
        return text, source

    return "", "все методы извлечения DOC не сработали"


# =========================
# RTF
# =========================

def read_rtf(file_path: str | Path) -> tuple[str, str | None]:
    """Читает RTF через striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        return "", "striprtf не установлен"

    path = Path(file_path)
    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл пустой или не найден"

    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            content = path.read_text(encoding=encoding, errors="replace")
            if content.strip():
                text = rtf_to_text(content).strip()
                if text:
                    return text, f"striprtf ({encoding})"
        except (OSError, UnicodeDecodeError):
            continue

    return "", "не удалось прочитать RTF"


# =========================
# Универсальное чтение
# =========================

def read_file(file_path: str | Path) -> tuple[str, str | None]:
    """Универсальное чтение файла по расширению."""
    path = Path(file_path)

    if not path.exists() or not path.is_file() or path.stat().st_size == 0:
        return "", "файл не существует или пустой"

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return "", f"неподдерживаемый формат: {path.suffix}"

    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".doc":
        return read_doc(path)
    if suffix == ".rtf":
        return read_rtf(path)

    return "", "неизвестный формат"


# =========================
# Структура результата
# =========================

@dataclass(slots=True)
class ParsedDocument:
    """Структура обработанного документа."""
    doc_name: str
    filepath: str
    filetype: str
    text: str
    chunks: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Парсер документов с разбиением текста на фрагменты."""

    def __init__(
        self,
        chunk_size: int = 1200,
        chunk_overlap: int = 200,
        min_chunk_size: int = 120,
    ) -> None:
        self.chunk_size = max(300, chunk_size)
        self.chunk_overlap = max(0, min(chunk_overlap, self.chunk_size // 2))
        self.min_chunk_size = max(50, min_chunk_size)

    @staticmethod
    def normalize_text(text: str) -> str:
        """Нормализует текст."""
        if not text:
            return ""
        text = text.replace("\x00", " ")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def split_paragraphs(self, text: str) -> list[str]:
        """Разбивает текст на логические параграфы."""
        parts = re.split(r"\n\s*\n", text)
        result: list[str] = []

        for part in parts:
            part = part.strip()
            if not part:
                continue
            if len(part) > int(self.chunk_size * 1.5):
                result.extend(self.split_long_text(part))
            else:
                result.append(part)

        return result

    def split_long_text(self, text: str) -> list[str]:
        """Разбивает длинный текст по предложениям."""
        sentences = re.split(r"(?<=[.!?])\s+", text)
        if len(sentences) < 2:
            return self.hard_split(text)

        result: list[str] = []
        current = ""

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            candidate = f"{current} {sentence}".strip() if current else sentence
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    result.append(current.strip())
                    overlap = current[-self.chunk_overlap:] if self.chunk_overlap else ""
                    overlap = self.smart_overlap(overlap)
                    current = f"{overlap} {sentence}".strip() if overlap else sentence
                else:
                    result.extend(self.hard_split(sentence))
                    current = ""

        if current.strip():
            result.append(current.strip())

        return result

    def hard_split(self, text: str) -> list[str]:
        """Принудительно разбивает длинный текст."""
        parts: list[str] = []
        start = 0

        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            if end < len(text):
                split_pos = text.rfind(" ", start, end)
                if split_pos > start + self.min_chunk_size:
                    end = split_pos

            part = text[start:end].strip()
            if part:
                parts.append(part)

            if end >= len(text):
                break

            start = max(end - self.chunk_overlap, start + 1)

        return parts

    @staticmethod
    def smart_overlap(text: str) -> str:
        """Подчищает overlap, чтобы не начинать с обрывка слова."""
        text = text.strip()
        if not text:
            return ""
        split_pos = text.find(" ")
        if 0 < split_pos < len(text) // 2:
            text = text[split_pos + 1:].strip()
        return text

    @staticmethod
    def extract_formulas(text: str) -> list[dict[str, Any]]:
        """Пытается извлечь простые формулы."""
        formulas: list[dict[str, Any]] = []
        patterns = [
            r"[A-Za-zА-Яа-я0-9_]+\s*=\s*[^=\n]{3,120}",
            r"\bQ\s*=\s*[^=\n]{3,120}",
            r"\bR\s*=\s*[^=\n]{3,120}",
        ]

        for pattern in patterns:
            for match in re.findall(pattern, text):
                raw = match.strip()
                if len(raw) < 4:
                    continue
                formulas.append(
                    {
                        "raw": raw,
                        "variables": sorted(set(re.findall(r"[A-Za-zА-Яа-я_]+", raw))),
                        "has_operator": any(op in raw for op in ("=", "+", "-", "*", "/")),
                    }
                )

        unique: list[dict[str, Any]] = []
        seen: set[str] = set()
        for item in formulas:
            key = item["raw"]
            if key not in seen:
                seen.add(key)
                unique.append(item)

        return unique[:20]

    @staticmethod
    def detect_table_like_content(text: str) -> bool:
        """Определяет, похож ли фрагмент на таблицу."""
        if "|" in text:
            return True

        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) < 2:
            return False

        tabular_lines = 0
        for line in lines[:12]:
            if re.search(r"\s{2,}", line):
                tabular_lines += 1
            elif len(re.findall(r"\d+", line)) >= 3:
                tabular_lines += 1

        return tabular_lines >= 2

    @staticmethod
    def build_metadata(
        path: Path,
        text: str,
        chunks: list[dict[str, Any]],
        source: str | None = None,
    ) -> dict[str, Any]:
        """Собирает метаданные документа."""
        return {
            "parsed": True,
            "filename": path.name,
            "file_stem": path.stem,
            "suffix": path.suffix.lower(),
            "size_bytes": path.stat().st_size,
            "char_count": len(text),
            "word_count": len(text.split()),
            "chunk_count": len(chunks),
            "has_formulas": any(chunk.get("has_formula", False) for chunk in chunks),
            "has_table_like_content": any(
                chunk.get("has_table_like_content", False) for chunk in chunks
            ),
            "extraction_source": source or "unknown",
            "platform": sys.platform,
        }

    def parse_file(self, file_path: str | Path) -> dict[str, Any]:
        """Парсит один файл."""
        path = Path(file_path)
        safe_print(f"📄 Парсинг: {path.name}")

        try:
            text, source = read_file(path)
            text = self.normalize_text(text)
        except (OSError, UnicodeDecodeError, ValueError, TypeError) as exc:
            safe_print(f"⚠️ Ошибка парсинга {path.name}: {exc}")
            return {
                "doc_name": path.name,
                "filepath": str(path),
                "filetype": path.suffix.lower(),
                "text": "",
                "chunks": [],
                "metadata": {
                    "parsed": False,
                    "filename": path.name,
                    "suffix": path.suffix.lower(),
                    "error": str(exc),
                },
            }

        if not text:
            reason = source or "неизвестная причина"
            safe_print(f"⚠️ Не удалось извлечь текст из {path.name}: {reason}")
            return {
                "doc_name": path.name,
                "filepath": str(path),
                "filetype": path.suffix.lower(),
                "text": "",
                "chunks": [],
                "metadata": {
                    "parsed": False,
                    "filename": path.name,
                    "suffix": path.suffix.lower(),
                    "reason": reason,
                },
            }

        safe_print(f"✅ Текст извлечён из {path.name}: {len(text)} символов через {source}")

        raw_chunks = self.split_paragraphs(text)
        chunks: list[dict[str, Any]] = []

        for idx, chunk_text in enumerate(raw_chunks):
            formulas = self.extract_formulas(chunk_text)
            chunks.append(
                {
                    "doc_name": path.name,
                    "filepath": str(path),
                    "chunk_id": idx,
                    "text": chunk_text,
                    "metadata": {"formulas": formulas},
                    "has_formula": bool(formulas),
                    "has_table_like_content": self.detect_table_like_content(chunk_text),
                }
            )

        metadata = self.build_metadata(path, text, chunks, source)

        return {
            "doc_name": path.name,
            "filepath": str(path),
            "filetype": path.suffix.lower(),
            "text": text,
            "chunks": chunks,
            "metadata": metadata,
        }

    def parse_directory(self, directory: str | Path, recursive: bool = True) -> list[dict[str, Any]]:
        """Парсит директорию с документами."""
        base = Path(directory)
        if not base.exists() or not base.is_dir():
            safe_print(f"⚠️ Директория не существует: {base}")
            return []

        pattern = "**/*" if recursive else "*"
        all_files = [p for p in base.glob(pattern) if p.is_file()]
        supported_files = [p for p in all_files if is_supported_file(p)]
        unsupported_files = [p for p in all_files if not is_supported_file(p)]

        safe_print(f"📦 Всего файлов в {base}: {len(all_files)}")
        safe_print(f"✅ Поддерживаемых: {len(supported_files)}")
        safe_print(f"⛔ Неподдерживаемых: {len(unsupported_files)}")

        if unsupported_files:
            safe_print("Неподдерживаемые файлы (первые 10):")
            for p in unsupported_files[:10]:
                safe_print(f"   - {p.name} [{p.suffix}]")

        parsed: list[dict[str, Any]] = []
        failed: list[str] = []

        for path in sorted(supported_files):
            try:
                item = self.parse_file(path)
                if item.get("text"):
                    parsed.append(item)
                else:
                    failed.append(path.name)
            except (OSError, UnicodeDecodeError, ValueError, TypeError) as exc:
                safe_print(f"⚠️ Критическая ошибка на файле {path.name}: {exc}")
                failed.append(path.name)

        safe_print(f"✅ Успешно распарсено: {len(parsed)}")
        if failed:
            safe_print(f"⚠️ Не удалось распарсить: {len(failed)}")
            for name in failed[:10]:
                safe_print(f"   - {name}")

        return parsed


def parse_file(file_path: str | Path, **kwargs: Any) -> dict[str, Any]:
    """Удобная функция парсинга одного файла."""
    parser = DocumentParser(**kwargs)
    return parser.parse_file(file_path)


def parse_directory(directory: str | Path, **kwargs: Any) -> list[dict[str, Any]]:
    """Удобная функция парсинга директории."""
    parser = DocumentParser(
        chunk_size=kwargs.pop("chunk_size", 1200),
        chunk_overlap=kwargs.pop("chunk_overlap", 200),
        min_chunk_size=kwargs.pop("min_chunk_size", 120),
    )
    recursive = kwargs.pop("recursive", True)
    return parser.parse_directory(directory, recursive=recursive)