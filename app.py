# -*- coding: utf-8 -*-
"""
Инженерный чат-бот для работы с документацией.
С ПРИНУДИТЕЛЬНЫМ ПОСТРОЕНИЕМ ИНДЕКСА И ПОДДЕРЖКОЙ OLLAMA.
"""

from __future__ import annotations

import asyncio
import json
import os
import time
from datetime import datetime

import streamlit as st
from reportlab.lib.colors import HexColor

try:
    from huggingface_hub import snapshot_download
except ImportError:
    snapshot_download = None
    print("⚠️ huggingface-hub не установлен. pip install huggingface-hub")

from core.agent_loop import AgentLoop
from core.config import PROCESSED_DIR, RAW_DIR, HF_DATASET_REPO_ID
from core.error_handler import ErrorHandler
from core.formula_engine import FormulaEngine
from core.parser import parse_directory
from core.prompts import get_quick_definition
from core.qa_engine import QASystem
from core.table_calculator import patch_app_with_table_calculator

st.set_page_config(
    page_title="Инженерный чат-бот",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded",
)

HISTORY_FILE = PROCESSED_DIR / "chat_history.json"
INDEX_FILE = PROCESSED_DIR / "faiss_index.pkl"


# ========= ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ =========

def run_async_safely(async_func, *args, **kwargs):
    """Безопасный запуск асинхронной функции в Streamlit."""
    loop = None
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        return loop.run_until_complete(async_func(*args, **kwargs))
    finally:
        if loop is not None:
            loop.close()
        asyncio.set_event_loop(None)


def call_maybe_async(func, *args, **kwargs):
    """Универсальный безопасный вызов sync/async функции."""
    if func is None:
        raise ValueError("Передана пустая функция (None) в call_maybe_async")

    if not callable(func):
        raise TypeError(f"Объект {type(func).__name__} не является вызываемым")

    result = func(*args, **kwargs)
    if asyncio.iscoroutine(result):
        return run_async_safely(lambda: result)
    return result


def get_initial_message() -> list[dict[str, str]]:
    """Начальное приветственное сообщение."""
    return [{
        "role": "assistant",
        "content": """🏗️ **Здравствуйте!** Я инженерный помощник по строительной документации.

📖 **База знаний:** ГОСТы, СП, технические регламенты и методические документы по строительству

**Что я умею:**
• 📖 Отвечать на вопросы по нормативной документации
• 📐 Рассчитывать толщину изоляции и теплопотери
• 🌍 Вычислять ГСОП (градусо-сутки отопительного периода)
• 💨 Определять расход теплоты на вентиляцию
• 📊 Находить таблицы и формулы в документах
• 🔍 Искать определения терминов

**Задайте свой вопрос или попросите сделать расчёт!**"""
    }]


def save_history() -> None:
    """Сохраняет историю чата в JSON."""
    HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(st.session_state.messages, f, ensure_ascii=False, indent=2)


# ========= СИНХРОНИЗАЦИЯ ДАТАСЕТА =========

def sync_hf_dataset_to_raw(force: bool = False) -> bool:
    """Скачивает документы из Hugging Face Dataset repo в RAW_DIR."""
    dataset_repo_id = (HF_DATASET_REPO_ID or "").strip()
    if not dataset_repo_id:
        print("ℹ️ HF_DATASET_REPO_ID не задан, синхронизация dataset пропущена")
        return False

    if snapshot_download is None:
        print("❌ huggingface-hub не установлен. Установите: pip install huggingface-hub")
        return False

    RAW_DIR.mkdir(parents=True, exist_ok=True)

    existing_docs = (
            list(RAW_DIR.glob("*.docx"))
            + list(RAW_DIR.glob("*.pdf"))
            + list(RAW_DIR.glob("*.rtf"))
            + list(RAW_DIR.glob("*.doc"))
    )

    if existing_docs and not force:
        print(f"✅ Документы уже есть в {RAW_DIR}: {len(existing_docs)} шт.")
        return True

    try:
        print(f"📥 Скачиваю dataset {dataset_repo_id} в {RAW_DIR} ...")
        snapshot_download(
            repo_id=dataset_repo_id,
            repo_type="dataset",
            local_dir=str(RAW_DIR),
        )

        downloaded_docs = (
                list(RAW_DIR.glob("*.docx"))
                + list(RAW_DIR.glob("*.pdf"))
                + list(RAW_DIR.glob("*.rtf"))
                + list(RAW_DIR.glob("*.doc"))
        )

        print(f"✅ Dataset синхронизирован. Найдено документов: {len(downloaded_docs)}")
        return bool(downloaded_docs)

    except Exception as dataset_error:
        print(f"⚠️ Ошибка загрузки dataset из Hugging Face: {dataset_error}")
        return False


# ========= ЭКСПОРТ =========

def export_history_to_docx():
    """Экспорт истории чата в DOCX."""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Инженерный чат-бот — история", 0)
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph()

        for msg in st.session_state.messages:
            role = "Пользователь" if msg["role"] == "user" else "Ассистент"
            doc.add_heading(role, level=1)
            doc.add_paragraph(msg["content"])
            doc.add_paragraph()

        output_path = PROCESSED_DIR / f"chat_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    except ImportError:
        st.error("❌ Для экспорта истории нужен python-docx: pip install python-docx")
        return None
    except (OSError, ValueError) as e:
        st.error(f"❌ Ошибка: {e}")
        return None


def export_to_docx(answer: str, sources: list, tables: list = None, formulas: list = None, filename: str = None):
    """Экспорт отчёта в DOCX."""
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"engineering_report_{timestamp}.docx"

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("Инженерный отчёт", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph()

        doc.add_heading("Ответ", level=1)
        doc.add_paragraph(answer)

        if tables:
            doc.add_heading("Таблицы", level=1)
            for table in tables[:2]:
                if isinstance(table, dict):
                    doc.add_paragraph(table.get("title", "Таблица"))
                    if table.get("content"):
                        doc.add_paragraph(table.get("content", ""))
                    elif table.get("rows"):
                        for row in table.get("rows", [])[:10]:
                            doc.add_paragraph(" | ".join(map(str, row)))
                    doc.add_paragraph()

        if formulas:
            doc.add_heading("Формулы", level=1)
            for formula in formulas[:3]:
                if isinstance(formula, dict):
                    raw = formula.get("raw") or formula.get("expression") or formula.get("name", "")
                    doc.add_paragraph(raw)
                    if formula.get("variables"):
                        doc.add_paragraph(f"Переменные: {', '.join(formula['variables'][:5])}")
                else:
                    doc.add_paragraph(str(formula))
                doc.add_paragraph()

        if sources:
            doc.add_heading("Источники", level=1)
            for src in sources:
                if isinstance(src, dict):
                    doc.add_paragraph(src.get("doc_name", "Документ"), style="List Bullet")
                else:
                    doc.add_paragraph(str(src), style="List Bullet")

        doc.add_paragraph()
        doc.add_paragraph("Отчёт сгенерирован автоматически.", style="Intense Quote")

        output_path = PROCESSED_DIR / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    except ImportError:
        st.error("❌ python-docx не установлен. pip install python-docx")
        return None
    except (OSError, ValueError) as e:
        st.error(f"❌ Ошибка создания DOCX: {e}")
        return None


def export_to_pdf(answer: str, sources: list, tables: list = None, formulas: list = None, filename: str = None):
    """Экспорт отчёта в PDF."""
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"engineering_report_{timestamp}.pdf"

    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        output_path = PROCESSED_DIR / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "TitleStyle",
            parent=styles["Title"],
            fontSize=24,
            textColor=HexColor("#1a5276"),
            alignment=TA_CENTER,
            spaceAfter=20,
        )

        heading_style = ParagraphStyle(
            "HeadingStyle",
            parent=styles["Heading1"],
            fontSize=16,
            textColor=HexColor("#2e86c1"),
            spaceAfter=12,
            spaceBefore=12,
        )

        normal_style = ParagraphStyle(
            "NormalStyle",
            parent=styles["Normal"],
            fontSize=11,
            spaceAfter=6,
        )

        story = [
            Paragraph("Инженерный отчёт", title_style),
            Spacer(1, 0.2 * inch),
            Paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style),
            Spacer(1, 0.2 * inch),
            Paragraph("Ответ", heading_style),
        ]

        for line in answer.split("\n"):
            if line.strip():
                clean_line = line.replace("**", "").replace("*", "")
                story.append(Paragraph(clean_line, normal_style))

        story.append(Spacer(1, 0.2 * inch))

        if sources:
            story.append(Paragraph("Источники", heading_style))
            for src in sources:
                if isinstance(src, dict):
                    story.append(Paragraph(f"• {src.get('doc_name', 'Документ')}", normal_style))
                else:
                    story.append(Paragraph(f"• {src}", normal_style))

        doc.build(story)
        return output_path

    except ImportError:
        st.warning("⚠️ reportlab не установлен. Будет создан DOCX вместо PDF.")
        return export_to_docx(answer, sources, tables, formulas, filename.replace(".pdf", ".docx"))
    except (OSError, ValueError) as e:
        st.error(f"❌ Ошибка создания PDF: {e}")
        return None


def render_export_buttons(
        answer: str,
        sources: list,
        tables: list,
        formulas: list,
        key_suffix: str = "current",
        response_id: int | None = None
):
    """Отображение кнопок экспорта с уникальными ключами."""
    if response_id is None:
        response_id = st.session_state.get("current_response_id", 0)

    unique_id = f"{key_suffix}_{response_id}_{int(time.time() * 1000)}"

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📄 Экспорт DOCX", key=f"export_docx_{unique_id}"):
            docx_path = export_to_docx(answer, sources, tables, formulas)
            if docx_path and docx_path.exists():
                with open(docx_path, "rb") as f:
                    st.download_button(
                        label="📥 Скачать DOCX",
                        data=f.read(),
                        file_name=docx_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{unique_id}",
                    )

    with col2:
        if st.button("📄 Экспорт PDF", key=f"export_pdf_{unique_id}"):
            pdf_path = export_to_pdf(answer, sources, tables, formulas)
            if pdf_path and pdf_path.exists():
                with open(pdf_path, "rb") as f:
                    st.download_button(
                        label="📥 Скачать PDF",
                        data=f.read(),
                        file_name=pdf_path.name,
                        mime="application/pdf",
                        key=f"download_pdf_{unique_id}",
                    )

    with col3:
        if st.button("📋 Копировать", key=f"copy_{unique_id}"):
            st.code(answer, language="text")
            st.success("✅ Текст скопирован!")


# ========= ИНИЦИАЛИЗАЦИЯ QA СИСТЕМЫ С OLLAMA =========

def force_rebuild_index(qa: QASystem) -> bool:
    """ПРИНУДИТЕЛЬНО перестраивает индекс с эмбеддингами."""
    print("=" * 50)
    print("🔨 ПРИНУДИТЕЛЬНАЯ ПЕРЕСТРОЙКА ИНДЕКСА")
    print("=" * 50)

    if not RAW_DIR.exists():
        print(f"❌ Папка {RAW_DIR} не существует")
        return False

    docs = list(RAW_DIR.glob("*.docx")) + list(RAW_DIR.glob("*.pdf")) + \
           list(RAW_DIR.glob("*.rtf")) + list(RAW_DIR.glob("*.doc"))

    print(f"📄 Найдено документов в RAW_DIR: {len(docs)}")

    if not docs:
        print("❌ Нет документов для индексации")
        return False

    print("📖 Начинаем парсинг документов...")
    parsed_docs = parse_directory(RAW_DIR, recursive=True)

    if not parsed_docs:
        print("❌ Парсинг не вернул ни одного документа")
        return False

    print(f"📄 Распарсено документов: {len(parsed_docs)}")

    total_chunks = sum(len(doc.get("chunks", [])) for doc in parsed_docs)
    print(f"🧩 Всего чанков: {total_chunks}")

    print("🔨 Строим индекс с эмбеддингами...")
    result = qa.build_index(parsed_docs)

    if not result:
        print("❌ build_index вернул False")
        return False

    print(f"✅ Индекс построен: {len(qa.chunks)} чанков")
    print(f"   embedding_model: {qa.embedding_model is not None}")
    print(f"   chunk_embeddings: {qa.chunk_embeddings is not None}")
    if qa.chunk_embeddings is not None:
        print(f"   embeddings shape: {qa.chunk_embeddings.shape}")

    print("💾 Сохраняем индекс...")
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    print(f"📁 Папка для сохранения: {PROCESSED_DIR}")
    print(f"📁 Права на запись: {os.access(PROCESSED_DIR, os.W_OK)}")

    save_result = qa.save_index(INDEX_FILE)

    if save_result:
        print(f"✅ Индекс сохранён: {INDEX_FILE}")
        if INDEX_FILE.exists():
            size = INDEX_FILE.stat().st_size
            print(f"   Размер файла: {size} bytes")
            if size < 1000000:
                print(f"⚠️ ВНИМАНИЕ: размер индекса слишком мал ({size} bytes)")
                return False
        else:
            print("❌ Файл не найден после сохранения!")
            return False
    else:
        print("❌ Ошибка сохранения индекса")
        return False

    print("=" * 50)
    return True


def get_llm_status(qa_system: QASystem) -> dict[str, str]:
    """Безопасно получает статус LLM."""
    status: dict[str, str] = {
        "use_llm": str(qa_system.use_llm),
        "llm_provider": qa_system.llm_provider,
        "llm_available": str(qa_system.llm_available),
        "selected_provider": "unknown",
        "ollama_alive": str(qa_system.is_ollama_alive()),
        "ollama_model": qa_system.ollama_model,
        "gemini_available": str(qa_system.gemini_available),
    }
    try:
        if hasattr(qa_system, 'get_selected_provider'):
            status["selected_provider"] = qa_system.get_selected_provider()
    except (AttributeError, TypeError, ValueError):
        pass
    return status


def init_qa_system() -> QASystem:
    """Инициализирует QA-систему с поддержкой Ollama."""
    # Определяем, какой LLM использовать
    use_llm = os.getenv("USE_LLM", "true").lower() == "true"
    llm_provider = os.getenv("LLM_PROVIDER", "ollama").strip().lower()
    ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434").strip()
    ollama_model = os.getenv("OLLAMA_MODEL", "phi3:mini").strip()
    gemini_api_key = os.getenv("GEMINI_API_KEY", "").strip()
    gemini_model = os.getenv("GEMINI_MODEL", "gemini-2.0-flash").strip()

    # Флаг принудительной перестройки индекса
    auto_sync = os.getenv("AUTO_SYNC_DATASET", "false").lower() == "true"
    auto_rebuild = os.getenv("AUTO_REBUILD_INDEX", "false").lower() == "true"

    print(f"🔧 Инициализация QASystem:")
    print(f"   use_llm: {use_llm}")
    print(f"   llm_provider: {llm_provider}")
    print(f"   ollama_base_url: {ollama_base_url}")
    print(f"   ollama_model: {ollama_model}")
    print(f"   auto_sync: {auto_sync}")
    print(f"   auto_rebuild: {auto_rebuild}")

    qa = QASystem(
        use_llm=use_llm,
        llm_provider=llm_provider if use_llm else "none",
        use_embeddings=True,
        ollama_base_url=ollama_base_url,
        ollama_model=ollama_model,
        gemini_api_key=gemini_api_key if use_llm and llm_provider == "gemini" else None,
        gemini_model=gemini_model,
    )

    # Загружаем индекс если есть
    if INDEX_FILE.exists():
        print(f"📂 Индекс найден: {INDEX_FILE}")
        try:
            qa.load_index(INDEX_FILE)
            print(f"✅ Индекс загружен: {len(qa.chunks)} чанков")
            return qa
        except (OSError, ValueError, TypeError) as e:
            print(f"⚠️ Ошибка загрузки индекса: {e}")
            print("🔄 Будет выполнена перестройка...")

    # Если индекс не загрузился
    if auto_sync:
        print("📥 AUTO_SYNC_DATASET=true → синхронизация dataset")
        sync_hf_dataset_to_raw(force=False)

    if auto_rebuild:
        print("🔨 AUTO_REBUILD_INDEX=true → перестройка индекса")
        force_rebuild_index(qa)
    else:
        print("⏭️ Автоперестройка индекса отключена")

    return qa


def init_session_state() -> None:
    """Инициализирует состояние сессии."""
    if "qa_system" not in st.session_state:
        with st.spinner("Загрузка системы..."):
            st.session_state.qa_system = init_qa_system()
            st.session_state.formula_engine = FormulaEngine(st.session_state.qa_system)
            st.session_state.agent_loop = AgentLoop(
                st.session_state.qa_system,
                st.session_state.formula_engine,
            )
            patch_app_with_table_calculator()

    if "error_handler" not in st.session_state:
        st.session_state.error_handler = ErrorHandler(log_level="info")

    if "messages" not in st.session_state:
        if HISTORY_FILE.exists():
            try:
                with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                    st.session_state.messages = json.load(f)
            except (json.JSONDecodeError, OSError):
                st.session_state.messages = get_initial_message()
        else:
            st.session_state.messages = get_initial_message()

    st.session_state.setdefault("current_answer", "")
    st.session_state.setdefault("current_sources", [])
    st.session_state.setdefault("current_tables", [])
    st.session_state.setdefault("current_formulas", [])
    st.session_state.setdefault("current_response_id", 0)


def auto_load_documents() -> bool:
    """Автоматическая загрузка и индексация документов."""
    qa_system = st.session_state.qa_system

    if qa_system.is_ready:
        chunks_count = len(qa_system.chunks) if hasattr(qa_system, 'chunks') else 0
        st.sidebar.success(f"✅ База знаний готова\n📄 {chunks_count} фрагментов")
        return True

    if not INDEX_FILE.exists():
        st.sidebar.info("🔄 Индекс отсутствует. Выполняется перестройка...")
        with st.spinner("📚 Индексация документов..."):
            if force_rebuild_index(qa_system):
                st.success("✅ Индекс перестроен")
                return True
            else:
                st.error("❌ Не удалось перестроить индекс")
                return False

    try:
        if qa_system.load_index(INDEX_FILE):
            st.sidebar.success(f"✅ Индекс загружен\n📄 {len(qa_system.chunks)} фрагментов")
            return True
    except (OSError, ValueError, TypeError) as e:
        st.sidebar.warning(f"⚠️ Ошибка загрузки индекса: {e}")

    return False


# ========= ОСНОВНОЙ ИНТЕРФЕЙС =========

def render_sidebar(
        qa_system: QASystem,
        formula_engine: FormulaEngine,
        error_handler: ErrorHandler
) -> None:
    """Рендер боковой панели."""
    with st.sidebar:
        st.header("📚 О системе")
        st.markdown("""
        - ✅ Семантический поиск по тексту
        - ✅ Инженерные расчёты
        - ✅ Извлечение нормативных параметров
        - ✅ Поиск таблиц и формул
        - ✅ Определения терминов
        """)
        st.divider()

        # ========= ДИАГНОСТИКА ИНДЕКСА =========
        st.subheader("🔍 Диагностика индекса")
        st.write(f"INDEX_FILE: `{INDEX_FILE}`")
        st.write(f"Файл существует: `{INDEX_FILE.exists()}`")
        if INDEX_FILE.exists():
            st.write(f"Размер: `{INDEX_FILE.stat().st_size}` bytes")

        chunks_count = len(qa_system.chunks) if hasattr(qa_system, 'chunks') else 0
        st.write(f"Чанков в памяти: `{chunks_count}`")
        st.write(f"is_ready: `{qa_system.is_ready}`")

        if qa_system.chunk_embeddings is not None:
            st.write(f"Эмбеддинги: `{qa_system.chunk_embeddings.shape}`")
        else:
            st.write("Эмбеддинги: `Нет`")

        st.divider()

        # ========= ДИАГНОСТИКА LLM =========
        st.subheader("🤖 Статус LLM")
        llm_status = get_llm_status(qa_system)
        st.write(f"use_llm: `{llm_status['use_llm']}`")
        st.write(f"llm_provider: `{llm_status['llm_provider']}`")
        st.write(f"llm_available: `{llm_status['llm_available']}`")
        st.write(f"selected_provider: `{llm_status['selected_provider']}`")
        st.write(f"ollama_alive: `{llm_status['ollama_alive']}`")
        st.write(f"ollama_model: `{llm_status['ollama_model']}`")
        st.write(f"gemini_available: `{llm_status['gemini_available']}`")
        st.divider()

        auto_load_documents()
        st.divider()

        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Перезагрузить индекс", use_container_width=True):
                if INDEX_FILE.exists():
                    qa_system.load_index(INDEX_FILE)
                    st.success(f"✅ Индекс перезагружен: {len(qa_system.chunks)} фрагментов")
                    st.rerun()
                else:
                    st.warning("⚠️ Индекс не найден")

        with col2:
            if st.button("🗑️ Очистить индекс", use_container_width=True):
                if INDEX_FILE.exists():
                    INDEX_FILE.unlink(missing_ok=True)
                    qa_system.is_ready = False
                    st.success("✅ Индекс очищен")
                    st.rerun()

        if st.button("📥 Синхронизировать dataset", use_container_width=True):
            with st.spinner("Скачивание документов из Hugging Face..."):
                if sync_hf_dataset_to_raw(force=True):
                    st.success("✅ Dataset синхронизирован")
                else:
                    st.error("❌ Не удалось синхронизировать dataset")
                st.rerun()

        if st.button("🔨 Перестроить индекс (с эмбеддингами)", use_container_width=True):
            with st.spinner("🔄 Перестройка индекса..."):
                if force_rebuild_index(qa_system):
                    st.success("✅ Индекс перестроен")
                    st.rerun()
                else:
                    st.error("❌ Ошибка перестройки индекса")

        if not qa_system.is_ready:
            if st.button("📚 Индексировать документы", key="index_btn", use_container_width=True):
                with st.spinner("Индексация..."):
                    if force_rebuild_index(qa_system):
                        st.success(f"✅ Проиндексировано {len(qa_system.chunks)} фрагментов")
                        st.rerun()
                    else:
                        st.error("❌ Не найдено документов для индексации")

        st.divider()

        st.subheader("📐 Доступные формулы")
        available_formulas = formula_engine.get_available_formulas()
        for formula in available_formulas:
            with st.expander(f"📖 {formula['name']}"):
                st.markdown(formula.get("expression", ""))
                st.caption(formula.get("description", ""))
                if formula.get("legend"):
                    st.markdown("**Обозначения:**")
                    st.markdown(formula["legend"])
                st.caption(f"📚 {formula.get('source', '')}")

        st.divider()

        st.subheader("📊 Статистика базы")
        docs_count = (
                len(list(RAW_DIR.glob("*.docx")))
                + len(list(RAW_DIR.glob("*.pdf")))
                + len(list(RAW_DIR.glob("*.rtf")))
                + len(list(RAW_DIR.glob("*.doc")))
        )
        chunks_count = len(qa_system.chunks) if qa_system.is_ready else 0

        col1, col2 = st.columns(2)
        col1.metric("Документов", docs_count)
        col2.metric("Фрагментов", chunks_count)

        st.divider()

        st.subheader("💾 Экспорт")
        if st.button("📄 Экспорт истории (DOCX)", use_container_width=True):
            docx_path = export_history_to_docx()
            if docx_path and docx_path.exists():
                with open(docx_path, "rb") as f:
                    st.download_button(
                        label="📥 Скачать DOCX",
                        data=f.read(),
                        file_name=docx_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
            else:
                st.error("❌ Ошибка создания файла")

        if st.button("🗑️ Очистить историю", use_container_width=True):
            st.session_state.messages = get_initial_message()
            if HISTORY_FILE.exists():
                HISTORY_FILE.unlink()
            st.session_state.current_answer = ""
            st.session_state.current_sources = []
            st.session_state.current_tables = []
            st.session_state.current_formulas = []
            st.rerun()

        if getattr(error_handler, "errors", None):
            st.divider()
            st.subheader("⚠️ Ошибки")
            with st.expander(f"Показать {len(error_handler.errors)} ошибок"):
                for i, err in enumerate(error_handler.errors[-5:], start=1):
                    st.error(f"{i}. {err.get('type', 'Error')}: {err.get('message', '')[:100]}")


def render_sources(sources: list) -> None:
    """Рендерит источники."""
    with st.expander("📚 Источники", expanded=False):
        for src in sources[:5]:
            if isinstance(src, dict):
                doc_name = src.get("doc_name") or src.get("docname") or src.get("source") or "Документ"
            else:
                doc_name = str(src)
            st.caption(f"📄 {doc_name}")


def render_tables(tables: list) -> None:
    """Рендерит таблицы."""
    with st.expander("📊 Таблицы", expanded=False):
        for table in tables[:3]:
            if isinstance(table, dict):
                rows = table.get("rows", [])
                if rows and isinstance(rows, list):
                    st.write(f"**{table.get('title', 'Таблица')}**")
                    st.dataframe(rows, use_container_width=True)
                else:
                    st.write(table)
            else:
                st.write(table)


def render_formulas(formulas: list) -> None:
    """Рендерит формулы."""
    with st.expander("📐 Формулы", expanded=False):
        for formula in formulas[:3]:
            if isinstance(formula, dict):
                st.code(formula.get("raw", ""))
            else:
                st.code(str(formula))


def render_reasoning(steps: list) -> None:
    """Рендерит цепочку рассуждений."""
    with st.expander("🧠 Цепочка рассуждений", expanded=False):
        for step in steps:
            st.caption(
                f"Шаг {step['step']}: {step['description']} "
                f"(уверенность: {step['confidence']:.0%})"
            )


def main() -> None:
    """Основная функция приложения."""
    init_session_state()

    qa_system = st.session_state.qa_system
    formula_engine = st.session_state.formula_engine
    agent_loop = st.session_state.agent_loop
    error_handler = st.session_state.error_handler

    st.title("🏗️ Инженерный помощник проектировщика")
    st.caption("📄 База знаний: ГОСТы, СП, технические регламенты и методические документы по строительству")

    render_sidebar(qa_system, formula_engine, error_handler)

    # Отрисовка истории сообщений
    for i, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

            if msg["role"] == "assistant":
                has_sources = bool(msg.get("sources"))
                has_tables = bool(msg.get("tables"))
                has_formulas = bool(msg.get("formulas"))

                if has_sources or has_tables or has_formulas:
                    with st.expander("📎 Источники и материалы", expanded=False):
                        if has_sources:
                            st.markdown("**Источники:**")
                            for src in msg.get("sources", []):
                                if isinstance(src, dict):
                                    st.markdown(f"- {src.get('doc_name', 'Документ')}")
                                else:
                                    st.markdown(f"- {src}")

                        if has_tables:
                            st.markdown("**Таблицы:**")
                            for table in msg.get("tables", [])[:2]:
                                if isinstance(table, dict):
                                    st.markdown(f"- {table.get('title', 'Таблица')}")

                        if has_formulas:
                            st.markdown("**Формулы:**")
                            for formula in msg.get("formulas", [])[:5]:
                                if isinstance(formula, dict):
                                    raw = formula.get("raw") or formula.get("expression") or formula.get("name") or "Формула"
                                    st.markdown(f"- `{raw}`")
                                else:
                                    st.markdown(f"- `{formula}`")

                    render_export_buttons(
                        answer=msg["content"],
                        sources=msg.get("sources", []),
                        tables=msg.get("tables", []),
                        formulas=msg.get("formulas", []),
                        key_suffix=f"history_{i}",
                        response_id=i,
                    )

    prompt = st.chat_input("Задайте вопрос по строительной документации...", key="main_chat_input")
    if not prompt:
        return

    st.session_state.messages.append({"role": "user", "content": prompt})

    with st.chat_message("user"):
        st.markdown(prompt)

    response = "Не удалось сформировать ответ."
    sources: list = []
    tables: list = []
    formulas: list = []

    with st.chat_message("assistant"):
        with st.spinner("🔍 Анализирую запрос..."):
            try:
                prompt_clean = prompt.strip()
                prompt_lower = prompt_clean.lower()

                # =========================================================
                # ТРИГГЕРЫ — только в НАЧАЛЕ запроса
                # =========================================================
                calc_triggers = [
                    "рассчитай", "вычисли", "посчитай", "толщин", "температур",
                    "потери", "формул", "вентиляц", "расход", "гсоп", "градусо"
                ]

                definition_triggers = [
                    "что такое ",
                    "что значит ",
                    "что означает ",
                    "что это ",
                    "дай определение ",
                    "дайте определение ",
                    "определение ",
                    "определи ",
                    "термин ",
                    "понятие ",
                    "расшифруй ",
                    "расшифровка ",
                    "аббревиатура ",
                ]

                table_triggers = [
                    "таблица",
                    "таблицы",
                    "таблицу",
                    "таблиц",
                    "табл",
                    "покажи таблиц",
                    "выведи таблиц",
                ]

                is_calc = any(w in prompt_lower for w in calc_triggers)
                is_definition_query = any(prompt_lower.startswith(t) for t in definition_triggers)
                is_table = any(w in prompt_lower for w in table_triggers)

                # =========================================================
                # 1. ОПРЕДЕЛЕНИЯ — ТОЛЬКО ПО ЯВНЫМ ТРИГГЕРАМ
                # =========================================================
                if is_definition_query:
                    clean_term = prompt_lower
                    for trigger in definition_triggers:
                        if clean_term.startswith(trigger):
                            clean_term = clean_term[len(trigger):].strip(" ?!.,:;\"'«»()[]")
                            break

                    if clean_term:
                        quick_def = get_quick_definition(clean_term)
                    else:
                        quick_def = None

                    if quick_def:
                        response = (
                            f"📖 **Определение:**\n\n"
                            f"{quick_def.get('definition', '')}\n\n"
                            f"📚 **Источник:** {quick_def.get('source', '')}"
                        )
                        if quick_def.get("example"):
                            response += f"\n\n📌 **Пример:** {quick_def['example']}"
                    elif clean_term:
                        definition_result = qa_system.find_definition(clean_term)
                        if definition_result.get("found"):
                            response = (
                                f"📖 **Определение термина «{clean_term}»:**\n\n"
                                f"{definition_result.get('definition', '')}\n\n"
                                f"📚 **Источник:** {definition_result.get('source', 'Нормативная база')}"
                            )
                        else:
                            response = f"⚠️ В загруженных документах не найдено определение для термина «{clean_term}»."
                    else:
                        response = "⚠️ Уточните термин для определения."

                # =========================================================
                # 2. РАСЧЁТЫ
                # =========================================================
                elif is_calc:
                    result = call_maybe_async(formula_engine.answer_calculation, prompt_clean)
                    response = result.get("answer", "Не удалось выполнить расчёт")
                    sources = result.get("sources", [])
                    tables = result.get("tables", [])
                    formulas = result.get("formulas", [])

                    if not formulas and result.get("formula"):
                        formulas = [result["formula"]]

                # =========================================================
                # 3. ТАБЛИЦЫ
                # =========================================================
                elif is_table:
                    result = qa_system.answer(prompt_clean)
                    response = result.get("answer", "Таблица не найдена")
                    tables = result.get("tables", [])
                    sources = result.get("sources", [])
                    formulas = result.get("formulas", [])

                    if tables:
                        response += "\n\n📊 **Найденные таблицы:**\n"
                        for table in tables[:2]:
                            if isinstance(table, dict):
                                response += f"\n**{table.get('title', 'Таблица')}**\n"
                                content = table.get("content", "")
                                if len(content) > 500:
                                    content = content[:500] + "..."
                                response += f"```\n{content}\n```\n"

                # =========================================================
                # 4. ВСЁ ОСТАЛЬНОЕ — AGENT LOOP
                # =========================================================
                else:
                    result = call_maybe_async(agent_loop.run, prompt_clean)
                    response = result.get("answer", "Не удалось получить ответ")
                    sources = result.get("sources", [])
                    tables = result.get("tables", [])
                    formulas = result.get("formulas", [])

                    if result.get("needs_clarification"):
                        questions = result.get("questions", [])
                        if questions:
                            response += "\n\n❓ **Уточните:**\n" + "\n".join([f"• {q}" for q in questions])

            except Exception as e:
                error_info = error_handler.handle(e, {"query": prompt})
                response = error_info.get("user_message", f"❌ Ошибка: {e}")

        # =========================================================
        # ОТОБРАЖЕНИЕ ОТВЕТА
        # =========================================================
        st.markdown(response)

        has_sources = bool(sources)
        has_tables = bool(tables)
        has_formulas = bool(formulas)

        if has_sources or has_tables or has_formulas:
            with st.expander("📎 Источники и материалы", expanded=False):
                if has_sources:
                    st.markdown("**Источники:**")
                    for src in sources:
                        if isinstance(src, dict):
                            st.markdown(f"- {src.get('doc_name', 'Документ')}")
                        else:
                            st.markdown(f"- {src}")

                if has_tables:
                    st.markdown("**Таблицы:**")
                    for table in tables[:5]:
                        if isinstance(table, dict):
                            st.markdown(f"- {table.get('title', 'Таблица')}")
                        else:
                            st.markdown(f"- {table}")

                if has_formulas:
                    st.markdown("**Формулы:**")
                    for formula in formulas[:5]:
                        if isinstance(formula, dict):
                            raw = formula.get("raw") or formula.get("expression") or formula.get("name") or "Формула"
                            st.markdown(f"- `{raw}`")
                        else:
                            st.markdown(f"- `{formula}`")

        current_response_id = st.session_state.get("current_response_id", 0) + 1
        st.session_state.current_response_id = current_response_id

        render_export_buttons(
            answer=response,
            sources=sources,
            tables=tables,
            formulas=formulas,
            key_suffix="current",
            response_id=current_response_id,
        )

    # Сохраняем сообщение в историю
    st.session_state.messages.append({
        "role": "assistant",
        "content": response,
        "sources": sources,
        "tables": tables,
        "formulas": formulas,
    })

    save_history()

if __name__ == "__main__":
    main()